"""Regenerate the synthetic demo sample files.

Run this once to refresh ``acme_q1_brief.txt``, ``acme_risk_factors.docx``
and ``acme_prospectus.pdf``. The generated files are committed so most
contributors never need to run the script. All content is fictional.

Usage (from the repo root, using the same Python used for the API):

    python docs/samples/_generate.py

This script uses only dependencies that are already in ``requirements.txt``
(python-docx). The PDF is built by hand so we avoid pulling in a new
dependency like reportlab.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

OUT_DIR = Path(__file__).resolve().parent


TXT_CONTENT = """Acme Corp — Q1 2026 Internal Diligence Brief (FICTIONAL SAMPLE)

Business overview
-----------------
Acme Corp is a fictional enterprise software vendor included here solely
as a demo sample. It sells a financial reporting suite ("AcmeClose") to
mid-market public companies in the United States and Canada. Primary
personas are the controller, the VP of FP&A, and the SEC reporting lead.
All names, numbers, and events below are invented.

Financial summary
-----------------
Q1 2026 revenue was $5.2M, up 12% year over year. Subscription revenue
was $4.6M (89% of total), services revenue was $0.6M. Gross margin
held at 71% on a GAAP basis and 76% on a non-GAAP basis. Operating
margin improved to 9% vs 4% in Q1 2025 on slower hiring and one-time
consulting credits.

Customer metrics
----------------
Ending ARR was $19.8M, up from $18.2M at year end. Net new ARR in
Q1 was $1.6M. Gross retention was 94%, net retention was 108%. The top
10 customers represented 22% of ARR; no single customer was over 5%.

Operations
----------
Headcount grew from 82 to 97 during the quarter, with engineering and
customer success accounting for 12 of the 15 net adds. The executive
team is unchanged. The company opened a secondary office in Raleigh, NC
on a three-year lease (~$18K/month) to support the customer success
build-out.

Notable items
-------------
1. Acme signed a new multi-year reseller agreement with "Lexton
   Accounting Group" in late Q1. Revenue impact is not expected until
   Q3 2026.
2. One material customer ("BlueHarbor Industrials") put their renewal
   on hold in February pending a platform migration decision. The
   $420K ARR is flagged as at-risk.
3. The company increased its cyber insurance coverage from $5M to $10M
   following a general industry recommendation from the auditor.

This brief was produced internally by the FP&A team and is not audited.
"""


RISK_FACTORS_HEADING = "Acme Corp — Selected Risk Factors (FICTIONAL SAMPLE)"

RISK_FACTORS_PARAGRAPHS = [
    "The following risk factors are fictional and are included solely to "
    "exercise the Deal Room AI summary and risk analysis flow. They do "
    "not describe a real company.",
    "Customer concentration. While no single customer accounted for more "
    "than 5% of ARR at the end of Q1 2026, the top ten customers "
    "represent roughly 22% of ARR. Loss of several of these customers in "
    "a short period could materially impact revenue and cash flow.",
    "Renewal risk with BlueHarbor Industrials. In February 2026, "
    "BlueHarbor placed its annual renewal on hold pending an internal "
    "platform migration decision. The associated $420K of ARR is "
    "considered at risk until a decision is reached.",
    "Litigation. Acme is a defendant in a commercial dispute with a "
    "former supplier alleging approximately $1.2M in unpaid services. "
    "The company has recorded no accrual and intends to defend the claim "
    "vigorously, but an adverse outcome could have a material impact in "
    "the quarter it occurs.",
    "Key person risk. The company depends heavily on its CTO and VP of "
    "Engineering for architectural decisions and the AcmeClose product "
    "roadmap. The company has not yet implemented formal succession "
    "plans for these roles.",
    "Cyber and data protection. AcmeClose processes sensitive financial "
    "reporting data. A breach, ransomware event, or extended service "
    "disruption could trigger customer contract penalties, reputational "
    "harm, and regulatory attention. The company increased its cyber "
    "insurance coverage in Q1 2026 from $5M to $10M.",
    "Macroeconomic exposure. Prolonged weakness in the mid-market public "
    "company segment, or renewed pressure on software budgets, could "
    "slow new ARR and weaken gross retention below the 94% observed in "
    "Q1 2026.",
    "Going concern. Management does not believe there is substantial "
    "doubt about the company's ability to continue as a going concern. "
    "Cash and equivalents at the end of Q1 2026 were $12.4M with no "
    "outstanding debt.",
]


PROSPECTUS_LINES = [
    "ACME CORP - Summary Prospectus (FICTIONAL SAMPLE)",
    "",
    "The following document is a fictional summary prospectus for",
    "Acme Corp, an invented software vendor. All facts, figures, and",
    "names are synthetic and should not be relied upon.",
    "",
    "Offering. Acme is offering up to 2,000,000 shares of Series C",
    "preferred stock at a target price of $25.00 per share, for",
    "gross proceeds of up to $50M. Proceeds are intended for general",
    "corporate purposes, including hiring in go-to-market, customer",
    "success expansion, and continued AcmeClose product investment.",
    "",
    "Company highlights. Ending ARR at March 31, 2026 was $19.8M with",
    "94% gross retention and 108% net retention. Q1 2026 revenue was",
    "$5.2M, up 12% year over year. Gross margin was 71% GAAP.",
    "",
    "Use of proceeds. Acme intends to use net proceeds approximately",
    "as follows: 45% engineering, 30% go-to-market, 15% customer",
    "success, and 10% general corporate purposes.",
    "",
    "Risk factors. See the accompanying Risk Factors document for",
    "a complete discussion of material risks, including customer",
    "concentration, litigation, renewal risk with BlueHarbor",
    "Industrials, and cyber exposure.",
    "",
    "This prospectus is fictional and is provided only to demonstrate",
    "the Deal Room AI document ingestion and citation flow.",
]


def write_txt() -> Path:
    path = OUT_DIR / "acme_q1_brief.txt"
    path.write_text(TXT_CONTENT, encoding="utf-8")
    return path


def write_docx() -> Path:
    path = OUT_DIR / "acme_risk_factors.docx"
    doc = Document()
    doc.add_heading(RISK_FACTORS_HEADING, level=1)
    for paragraph in RISK_FACTORS_PARAGRAPHS:
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return path


def _pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_pdf() -> Path:
    """Build a tiny text-only PDF by hand.

    The structure is intentionally minimal: Catalog, Pages, one Page, a
    Helvetica font, and a single content stream with a stack of text
    lines. pypdf can extract the text without any special handling.
    """
    font_size = 11
    leading = 14
    x, y0 = 72, 760  # left margin, top baseline

    content_lines: list[str] = ["BT", "/F1 " + str(font_size) + " Tf", f"{x} {y0} Td"]
    for i, line in enumerate(PROSPECTUS_LINES):
        escaped = _pdf_escape(line) if line else ""
        if i == 0:
            content_lines.append(f"({escaped}) Tj")
        else:
            content_lines.append(f"0 -{leading} Td")
            content_lines.append(f"({escaped}) Tj")
    content_lines.append("ET")
    content_stream = "\n".join(content_lines).encode("latin-1")

    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(
        b"<< /Type /Page /Parent 2 0 R "
        b"/Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>"
    )
    objects.append(
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>"
    )
    objects.append(
        (
            "<< /Length " + str(len(content_stream)) + " >>\nstream\n"
        ).encode("latin-1")
        + content_stream
        + b"\nendstream"
    )

    out = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets: list[int] = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode("latin-1")
        out += body
        out += b"\nendobj\n"

    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode("latin-1")
    out += (
        "trailer\n<< /Size "
        + str(len(objects) + 1)
        + " /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset)
        + "\n%%EOF\n"
    ).encode("latin-1")

    path = OUT_DIR / "acme_prospectus.pdf"
    path.write_bytes(bytes(out))
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    txt_path = write_txt()
    docx_path = write_docx()
    pdf_path = write_pdf()
    for p in (txt_path, docx_path, pdf_path):
        print(f"wrote {p.relative_to(OUT_DIR.parent.parent)} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
